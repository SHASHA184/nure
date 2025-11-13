"""
Script to create test files for Lab 3
"""

from docx import Document
from PIL import Image, ImageDraw, ImageFont
import os


def create_word_document():
    """Create a test Word document"""
    doc = Document()

    doc.add_heading('Тестовий документ для лабораторної роботи №3', 0)

    doc.add_heading('Вступ', level=1)
    doc.add_paragraph(
        'Цей документ створено для тестування власної хеш-функції. '
        'Він містить текст українською мовою та демонструє роботу '
        'з документами формату DOCX.'
    )

    doc.add_heading('Криптографічні хеш-функції', level=1)
    doc.add_paragraph(
        'Хеш-функція - це обчислювально ефективна функція, що відображає '
        'двійковий рядок довільної довжини в двійковий рядок деякої фіксованої довжини.'
    )

    doc.add_paragraph(
        'Основні властивості криптографічних хеш-функцій:'
    )

    doc.add_paragraph('Односпрямованість', style='List Bullet')
    doc.add_paragraph('Свобода від колізій', style='List Bullet')
    doc.add_paragraph('Повне перемішування', style='List Bullet')

    doc.add_heading('Висновок', level=1)
    doc.add_paragraph(
        'Даний документ буде використовуватись для обчислення '
        'хеш-значень та демонстрації колізій.'
    )

    # Save document
    doc.save('test_files/document.docx')
    print("Created: test_files/document.docx")


def create_test_image():
    """Create a test image"""
    # Create a colorful image with text
    width, height = 400, 300
    image = Image.new('RGB', (width, height), color='lightblue')
    draw = ImageDraw.Draw(image)

    # Draw some shapes
    draw.rectangle([50, 50, 150, 150], fill='red', outline='black', width=2)
    draw.ellipse([200, 50, 350, 200], fill='green', outline='black', width=2)
    draw.polygon([(200, 250), (300, 280), (250, 220)], fill='yellow', outline='black')

    # Add text
    try:
        # Try to use a default font, fallback to default if not available
        font = ImageFont.load_default()
    except:
        font = None

    draw.text((width // 2 - 50, 10), "Lab 3 Test Image", fill='black', font=font)
    draw.text((width // 2 - 80, height - 30), "Hash Function Testing", fill='black', font=font)

    # Save image
    image.save('test_files/image.png')
    print("Created: test_files/image.png")


if __name__ == "__main__":
    # Create test_files directory if it doesn't exist
    os.makedirs('test_files', exist_ok=True)

    # Create test files
    create_word_document()
    create_test_image()

    print("\nAll test files created successfully!")
